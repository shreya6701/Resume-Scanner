from PyPDF2 import PdfReader
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer

import docx
import math
import numpy as np

import nltk
nltk.download('stopwords')
nltk.download('punkt')
from nltk.corpus import stopwords

# import spacy
# nlp=spacy.load("en_core_web_sm")
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
cv= CountVectorizer()

import os

#word frequency of imp words
def word_countpdf(str):
    counts_pdf = dict()
    word_pdf = str

    for word in word_pdf:
        if word in counts_pdf:
            counts_pdf[word] += 1
        else:
            counts_pdf[word] = 1

    return counts_pdf

#weights of imp words
def check_word(word, file):
    final=[all([w in x for w in word]) for x in file]
    word_len=[file[i] for i in range(0, len(final)) if final[i]]
    return int(len(word_len))


def each_word(str):
    each_word_pdf = dict()
    word_pdf = str

    for word in word_pdf:
        if word in each_word_pdf:
            each_word_pdf[word] = check_word(word, word_pdf) 
        else:
            each_word_pdf[word] = 1

    return each_word_pdf


def get_keywords(text, top_n=20):
    vectorizer = TfidfVectorizer()
    tfidf = vectorizer.fit_transform([text])
    
    # Get top N keywords by TF-IDF score
    top_n_indices = np.argsort(tfidf.toarray())[0][::-1][:top_n]
    keywords = [vectorizer.get_feature_names_out()[i] for i in top_n_indices]
    
    return keywords

def get_similarity(resume, job_desc):
    content=[resume, job_desc]
    matrix=cv.fit_transform(content)

    similarity_matrix=cosine_similarity(matrix)
    print("Resume "+i+ " matches by : " + str(similarity_matrix[1][0]*100))

# calculate teh weighted sum
def get_weighted_score(resume_tokens, jd_keywords):
    score = 0
    for token in resume_tokens:
        if token in jd_keywords:
            score += jd_idf_score.get(token, 0)
    
    return score

def rank_resumes():
    
    ranked_resumes.append({'resume': i, 'similarity': similarity, 'weighted_score': weighted_score})
    
    return ranked_resumes

# # job description
# jd_document = docx.Document("JD_Network Support.docx")
# doc_text = '\n\n'.join(
#     paragraph.text for paragraph in jd_document.paragraphs
# )

# doc_text=doc_text.lower()

sw_nltk = stopwords.words('english')

# jd_words = [word for word in doc_text.split() if word.lower() not in sw_nltk]
# new_jdtext = " ".join(jd_words)
# # print("JD length: ", len(new_jdtext.split()))


# jd_imp_keywords = get_keywords(new_jdtext)

job_desc = input("Enter the important keywords from job descrption, separated by spaces: ")
jd_imp_keywords = job_desc.split()
print("Important keywords:", jd_imp_keywords)

jd_imp_keywords=[e.lower() for e in jd_imp_keywords]
# file2=nlp(new_jdtext)
# print(file2)

# imp_jd_words=(file2.ents)s
# print(imp_jd_words)


jd_tf_score=word_countpdf(jd_imp_keywords)
jd_tf_score.update((x,y/int(len(jd_imp_keywords))) for x, y in jd_tf_score.items())
# print(jd_tf_score)

jd_idf_score=each_word(jd_imp_keywords)
jd_idf_score.update((x,math.log(int(len(jd_imp_keywords))/y)) for x, y in jd_idf_score.items())
print(jd_idf_score)


# sum=0
# for i in jd_idf_score:
#     sum+=jd_idf_score.get(i,0)

# print(sum)
# print("\n")

# content=[new_jdtext, new_jdtext]

ranked_resumes = []

#resume
directory='resume'

for i in os.listdir(directory):
    try:
        # read the pdf file
        reader=PdfReader(os.path.join(directory,i))
        page=reader.pages[0]
        text=page.extract_text()
        # Total_words=len(text.split())
    
    except:
        # read the word document
        reader = docx.Document(os.path.join(directory,i))
        text = '\n\n'.join(
            paragraph.text for paragraph in reader.paragraphs
        )
        
        # print(len(docText.split()))

   
    text=text.lower()
    words = [word for word in text.split() if word.lower() not in sw_nltk]
    new_text = " ".join(words)
    len_new_text=len(new_text.split())
    # print("Resume " + i + " length: ", len_new_text)
 
    similarity = get_similarity(new_text, job_desc)
    print(similarity)

    resume_tokens = nltk.word_tokenize(new_text)
    # print (resume_tokens)

    weighted_score = get_weighted_score(resume_tokens, jd_imp_keywords)
    print(weighted_score)

    ranking=rank_resumes()
 

ranked_resumes = sorted(ranked_resumes, key=lambda x: x['weighted_score'], reverse=True)
print(ranked_resumes)

